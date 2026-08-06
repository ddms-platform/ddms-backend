# -*- coding: utf-8 -*-
"""Generate detailed DDMS project Word documentation."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(r"D:\PROJECT_CAPSTONE_2026\BE\ddms-backend\DDMS_Tai_Lieu_Chi_Tiet_Du_An.docx")


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size={1: 16, 2: 14, 3: 12}.get(level, 11), bold=True,
                     color=RGBColor(0x1F, 0x4E, 0x79) if level <= 2 else RGBColor(0x2E, 0x75, 0xB6))
    return h


def add_para(doc, text, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=11)
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # light gray background via shading
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    shd.set(qn("w:val"), "clear")
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # header bg
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F4E79")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=9)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # ========== COVER ==========
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TÀI LIỆU CHI TIẾT DỰ ÁN")
    set_run_font(r, size=22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("DDMS – Da Nang Dock Management System\n(Hệ thống Quản lý Du thuyền & Tour thuyền Đà Nẵng)")
    set_run_font(r, size=16, bold=True)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run(
        "\nCapstone Project 2026\n"
        "Phiên bản tài liệu: 1.0\n"
        "Ngày lập: 29/07/2026\n"
        "Phạm vi: Backend (.NET) + Frontend Customer/Owner + Frontend Admin/Kiosk"
    )
    set_run_font(r, size=12)

    doc.add_page_break()

    # ========== TOC-like overview ==========
    add_heading_styled(doc, "MỤC LỤC TÓM TẮT", 1)
    toc_items = [
        "1. Tổng quan dự án",
        "2. Kiến trúc hệ thống",
        "3. Công nghệ sử dụng (Tech Stack)",
        "4. Cơ sở dữ liệu & Entities",
        "5. Backend – Pipeline & Cấu hình",
        "6. Backend – Auth & Phân quyền (chi tiết hàm)",
        "7. Backend – Booking / Hold / Pay / Cancel / Check-in",
        "8. Backend – Billing PayOS & Wallet",
        "9. Backend – Owner Registration & Boat Compliance",
        "10. Backend – Admin Services",
        "11. Backend – Các Service còn lại (hàm + logic)",
        "12. Backend – Controllers & API Endpoints",
        "13. Backend – Background Jobs & SignalR Hubs",
        "14. Frontend Customer/Owner – Cấu trúc & Routes",
        "15. Frontend Customer – Auth, Axios, Token Refresh",
        "16. Frontend Customer – Services (hàm + endpoint)",
        "17. Frontend Customer – Pages & luồng UI",
        "18. Frontend Admin/Kiosk – Chi tiết",
        "19. Luồng nghiệp vụ end-to-end",
        "20. Envelope API, mã lỗi, ghi chú kỹ thuật",
    ]
    for item in toc_items:
        add_bullet(doc, item)

    doc.add_page_break()

    # ========== 1. TỔNG QUAN ==========
    add_heading_styled(doc, "1. TỔNG QUAN DỰ ÁN", 1)
    add_para(
        doc,
        "DDMS (Da Nang Dock Management System / Boat Tour Management System) là hệ thống quản lý "
        "du thuyền, lịch neo đậu, tour thuyền, đặt chỗ, thanh toán, kiểm soát giấy tờ pháp lý, "
        "chat realtime và quản trị vận hành tại khu vực Đà Nẵng (điểm neo tiêu biểu: Bến Du Thuyền Sông Hàn)."
    )
    add_para(doc, "Hệ thống gồm 3 ứng dụng chính:", bold=True)
    add_bullet(doc, "Backend API: D:\\PROJECT_CAPSTONE_2026\\BE\\ddms-backend\\DDMS.Backend (ASP.NET Core)")
    add_bullet(doc, "Frontend Customer + Owner: D:\\PROJECT_CAPSTONE_2026\\FE\\ddms-frontend (React + Vite, port 5173)")
    add_bullet(doc, "Frontend Admin + Kiosk: D:\\PROJECT_CAPSTONE_2026\\FE\\ddms-frontend-admin\\ddms-frontedn-admin (React + Vite, port 5174)")

    add_heading_styled(doc, "1.1. Các vai trò người dùng (Roles)", 2)
    add_table(
        doc,
        ["Role", "Mô tả", "Ứng dụng chính"],
        [
            ["user", "Khách hàng: tìm tour, đặt chỗ, ví, chat, wishlist, review", "ddms-frontend"],
            ["owner", "Chủ thuyền: quản lý thuyền/tour/lịch/giấy tờ/doanh thu", "ddms-frontend (+ một phần admin FE)"],
            ["admin", "Quản trị: duyệt owner, GCN, tour, bảo trì, rút tiền, docks", "ddms-frontend-admin"],
            ["agent", "Đại lý B2B: chính sách giữ chỗ dài hơn (hold policy)", "Backend logic BookingHold"],
        ],
    )

    add_heading_styled(doc, "1.2. Phạm vi chức năng chính", 2)
    for x in [
        "Đăng ký / đăng nhập (email + Google OAuth), xác thực email, quên mật khẩu, JWT + refresh token",
        "Đăng ký trở thành chủ thuyền (multipart hồ sơ + tàu + giấy tờ) → Admin duyệt",
        "Quản lý thuyền, cabin, dịch vụ kèm theo, ảnh (Cloudinary), bảo trì cảng",
        "Quản lý tour, lịch trình, route, FAQ, ảnh tour; tìm kiếm & catalog public",
        "Đặt chỗ: Create (pending) / Hold (holding + countdown) / Pay / Cancel / Check-in QR (kiosk)",
        "Ví điện tử (hoàn tiền hủy tour), rút tiền → Admin duyệt",
        "Billing chủ thuyền: hoa hồng + neo đậu + bảo trì → thanh toán PayOS + webhook + SignalR",
        "Tuân thủ giấy tờ thuyền (compliance job): Warning / Hidden / Locked",
        "Chat realtime SignalR giữa khách và chủ thuyền theo booking",
        "Thông báo in-app, email templates, wishlist, review",
        "Admin dashboard, duyệt promotions, docks, certificates, withdrawals, maintenances",
        "Kiosk check-in vé điện tử bằng mã QR / mã booking",
    ]:
        add_bullet(doc, x)

    # ========== 2. KIẾN TRÚC ==========
    add_heading_styled(doc, "2. KIẾN TRÚC HỆ THỐNG", 1)
    add_para(doc, "Kiến trúc tổng thể theo mô hình 3 lớp (layered) trên Backend và SPA trên Frontend:", bold=True)
    add_code(
        doc,
        "Browser (React SPA)\n"
        "   │  HTTPS / Axios  +  SignalR (/hub/chat, /hub/billing)\n"
        "   ▼\n"
        "ASP.NET Core API Controllers\n"
        "   │  JWT Auth + Role Authorize + Rate Limit (Auth)\n"
        "   ▼\n"
        "Services (Business Logic)\n"
        "   │\n"
        "   ▼\n"
        "Repositories + EF Core AppDbContext\n"
        "   │\n"
        "   ▼\n"
        "MySQL 8 (database: boat_tour)\n"
        "\n"
        "Tích hợp ngoài: PayOS, Cloudinary, SMTP Email, Google OAuth, Open-Meteo (FE weather)",
    )
    add_para(
        doc,
        "Luồng request Backend (Program.cs): Localization → GlobalExceptionMiddleware → Swagger → "
        "CORS → RateLimiter → Authentication → Authorization → MapHubs → MapControllers."
    )

    # ========== 3. TECH STACK ==========
    add_heading_styled(doc, "3. CÔNG NGHỆ SỬ DỤNG (TECH STACK)", 1)
    add_heading_styled(doc, "3.1. Backend", 2)
    add_table(
        doc,
        ["Thành phần", "Công nghệ"],
        [
            ["Framework", "ASP.NET Core (Web API)"],
            ["ORM", "Entity Framework Core + MySQL (Pomelo / MySqlServerVersion 8.0.21)"],
            ["Auth", "JWT Bearer, BCrypt password hash, Google ID Token validate"],
            ["Realtime", "SignalR (ChatHub, BillingHub)"],
            ["Validation", "FluentValidation AutoValidation"],
            ["Docs API", "Swashbuckle / Swagger"],
            ["Upload", "Cloudinary"],
            ["Payment", "PayOS SDK"],
            ["Background", "IHostedService / BackgroundService"],
            ["Email", "SMTP / EmailSender templates"],
        ],
    )
    add_heading_styled(doc, "3.2. Frontend Customer", 2)
    add_table(
        doc,
        ["Thành phần", "Công nghệ"],
        [
            ["UI", "React 19 + TypeScript + Vite 8"],
            ["Routing", "react-router-dom 7"],
            ["HTTP", "axios"],
            ["Realtime", "@microsoft/signalr"],
            ["State", "zustand + React Context (Auth/Theme/Language)"],
            ["UI kit", "Tailwind CSS 4, shadcn/radix, lucide, hugeicons"],
            ["i18n", "i18next (en/vn)"],
            ["Charts", "recharts"],
            ["Auth Google", "@react-oauth/google"],
            ["QR", "react-qr-code"],
        ],
    )
    add_heading_styled(doc, "3.3. Frontend Admin", 2)
    add_para(
        doc,
        "Tương tự customer (React 19, Vite, Tailwind, axios, i18n) nhưng không dùng SignalR/Google OAuth; "
        "có thêm html5-qrcode cho kiosk check-in. Default ngôn ngữ: tiếng Việt."
    )

    # ========== 4. DATABASE ==========
    add_heading_styled(doc, "4. CƠ SỞ DỮ LIỆU & ENTITIES", 1)
    add_para(
        doc,
        "Database MySQL tên boat_tour. Entities nằm trong Models/Entities. Các bảng chính:"
    )
    add_table(
        doc,
        ["Entity", "Trường / ý nghĩa chính"],
        [
            ["user", "id, full_name, email, password_hash?, phone, address, avatar_url, is_active, google_id, email_verified_at"],
            ["role / user_role", "admin | owner | user | agent"],
            ["refresh_token", "token_hash (SHA256), expires_at, revoked, IP, UA"],
            ["email_verification_token", "purpose EmailVerification/PasswordReset, hash, used_at"],
            ["owner_profile", "business_name, license, entity_type, status Pending/Verified/Rejected, is_verified"],
            ["owner_document", "document_type, document_url, expiry_date"],
            ["owner_payment", "amount, status pending/paid, payos_order_code"],
            ["boat", "owner_id, name, type, status, compliance_status, length/beam, registration_number, is_deleted"],
            ["boat_cabin", "capacity, price, total_rooms"],
            ["boat_service", "addon dịch vụ trên thuyền (price, is_active)"],
            ["boat_image", "image_url, public_id Cloudinary"],
            ["boat_certificate", "certificate_type, expiry_date, status Pending/Approved/Expired/Rejected"],
            ["certificate_type", "code, name_vi/en, scope boat|owner"],
            ["boat_maintenance", "liên kết dịch vụ cảng, status pending/approved"],
            ["tour", "name, price, duration, location, rating, status, cancel_policy"],
            ["tour_schedule", "tour_id, boat_id?, dock_id?, start/end, status"],
            ["tour_image / route / faq", "nội dung catalog tour"],
            ["dock / dock_schedule", "bến neo + lịch gán thuyền–bến"],
            ["booking", "giá breakdown, status, hold_expired_at, hold_reminder_sent, cancel_reason"],
            ["booking_cabin / booking_service", "quantity, unit_price"],
            ["promotion", "code, discount, usage, status pending/approved"],
            ["review", "rating, comment, media JSON, gắn booking"],
            ["wishlist", "user_id + tour_id"],
            ["user_wallet / wallet_withdrawal", "balance; rút tiền pending/approved/rejected"],
            ["notification + recipient", "title, body, is_read"],
            ["conversation / member / message", "chat theo booking"],
            ["audit_log", "table_name, action, old/new values"],
        ],
    )
    add_para(doc, "Views DB hỗ trợ báo cáo: v_dashboard, v_booking_detail, v_revenue_stat, v_top_tour, v_loyalty_balance, v_unread_notification, v_unread_message.")
    add_para(doc, "Trạng thái booking (BookingStatuses): pending | holding | confirmed | paid | completed | checked_in | cancelled.")
    add_para(doc, "Compliance boat: Valid | Warning | Hidden | Locked.")

    # ========== 5. BACKEND PIPELINE ==========
    add_heading_styled(doc, "5. BACKEND – PIPELINE & CẤU HÌNH", 1)
    add_heading_styled(doc, "5.1. Program.cs – Đăng ký dịch vụ", 2)
    for x in [
        "AddControllers + JSON camelCase",
        "AddDdmsLocalization, AddDdmsSwagger, AddRequestValidation (FluentValidation)",
        "AddProjectDependencies() – đăng ký hàng loạt Service/Repository scoped",
        "AddSignalR; PayOSClient singleton từ config PayOS",
        "Configure Options: Jwt, Cloudinary, Email, Google, Billing, BoatCompliance, BookingHold (ValidateOnStart)",
        "HostedService: BoatComplianceBackgroundService, SeatHoldCleanupBackgroundService",
        "DbContext MySQL; JWT Bearer + ConfigureDdmsJwtBearer (role claim = \"role\", MapInboundClaims=false)",
        "CORS theo AllowedOrigins; RateLimiter Auth 30 req/phút/IP",
        "Đăng ký thủ công thêm Auth/Boat/Dock/Chat/Notification repositories & services",
    ]:
        add_bullet(doc, x)

    add_heading_styled(doc, "5.2. Startup migrate + seed", 2)
    add_para(
        doc,
        "Khi khởi động: Database.Migrate(); dọn dock trùng tên \"Bến Du Thuyền Sông Hàn\"; "
        "re-seed dock_schedules cho mọi boat; đảm bảo role admin; seed user admin@ddms.com / Admin@123; "
        "gỡ admin role khỏi user không phải admin mặc định."
    )
    add_para(
        doc,
        "Lưu ý vận hành: seed/cleanup chạy mỗi lần start – cần tắt hoặc kiểm soát khi production.",
        bold=True,
    )

    add_heading_styled(doc, "5.3. GlobalExceptionMiddleware", 2)
    add_para(
        doc,
        "Bắt mọi exception chưa xử lý → trả ApiErrorResponse JSON. Map ValidationException→400, "
        "UnauthorizedException→401, ForbiddenException→403, NotFoundException→404, AppException theo ErrorCode "
        "(AuthInvalidCredentials→401, AuthEmailNotVerified→403, AuthAccountInactive→403, rate limit→429, "
        "refresh token errors→401…). Tour module errors được localize qua IErrorMessageLocalizer."
    )

    add_heading_styled(doc, "5.4. JwtBearerExtensions.ConfigureDdmsJwtBearer", 2)
    add_bullet(doc, "MapInboundClaims = false; RoleClaimType = \"role\"; NameClaimType = sub")
    add_bullet(doc, "OnChallenge: trả JSON 401 phân biệt token hết hạn (AuthTokenExpired) vs unauthorized")
    add_bullet(doc, "OnMessageReceived: đọc access_token từ query string khi path bắt đầu /hub (cho SignalR)")

    # ========== 6. AUTH ==========
    add_heading_styled(doc, "6. BACKEND – AUTH & PHÂN QUYỀN (CHI TIẾT HÀM)", 1)

    add_heading_styled(doc, "6.1. IAuthService / AuthService", 2)
    add_table(
        doc,
        ["Hàm", "Logic xử lý"],
        [
            ["RegisterAsync(RegisterRequest)", "Validate fullName/email/password (≥8, hoa, thường, số, đặc biệt); chuẩn hóa email; check trùng; tạo user BCrypt; gán role user; gửi link verify; Dev có thể trả verificationLink"],
            ["LoginAsync(LoginRequest, ip, ua)", "Validate; lấy user; từ chối nếu không password_hash hoặc sai MK; EnsureAccountActive + EnsureEmailVerified; IssueTokensAsync"],
            ["VerifyEmailAsync", "Ủy quyền EmailVerificationService.VerifyByTokenAsync; kiểm tra active; trả alreadyVerified"],
            ["ResendVerificationEmailAsync", "Rate-limit friendly; không leak email tồn tại; gửi lại link"],
            ["ForgotPasswordAsync", "PasswordResetService.SendResetLinkAsync"],
            ["ResetPasswordAsync", "ResetPasswordByTokenAsync + policy mật khẩu"],
            ["ChangePasswordAsync(userId, req)", "Verify current; new≠old; update hash; revoke ALL refresh tokens"],
            ["RefreshTokenAsync", "Hash raw token; invalid/expired; reuse detection nếu đã revoked → revoke all; rotate token mới"],
            ["LogoutAsync / LogoutAllAsync", "Revoke 1 token hoặc toàn bộ session"],
            ["GetMeAsync", "Profile + roles + hasOwnerProfile"],
            ["UpdateProfileAsync", "full_name, phone, address"],
            ["UpdateAvatarAsync", "Upload Cloudinary → avatar_url"],
        ],
    )

    add_heading_styled(doc, "6.2. ITokenService / TokenService", 2)
    add_bullet(doc, "GenerateAccessToken(user, roles): JWT claims sub, NameIdentifier, email, UniqueName, nhiều claim \"role\"; HMAC-SHA256; hết hạn accessTokenMinutes")
    add_bullet(doc, "GenerateRefreshToken(): 64 random bytes Base64")
    add_bullet(doc, "HashToken(raw): SHA256 hex – chỉ lưu hash vào DB")

    add_heading_styled(doc, "6.3. IAuthSessionService / AuthSessionService", 2)
    add_bullet(doc, "IssueTokensAsync: lấy roles từ user_roles; tạo access+refresh; lưu refresh_token (hash, IP, UA, expires); EnforceSessionLimitAsync (maxRefreshTokensPerUser); trả expiresInSeconds")
    add_bullet(doc, "EnsureEmailVerified / EnsureAccountActive: ném AppException nếu chưa verify hoặc inactive")

    add_heading_styled(doc, "6.4. IGoogleAuthService / GoogleAuthService.LoginWithGoogleAsync", 2)
    for step in [
        "Validate idToken + GoogleOptions.clientId",
        "GoogleJsonWebSignature.ValidateAsync (Audience, clock skew 5 phút)",
        "Tìm user theo google_id hoặc email",
        "Chưa có: tạo user mới (google_id, avatar, email_verified_at=now), role user",
        "Đã có: gắn google_id nếu thiếu; conflict → lỗi; auto verify; cập nhật avatar nếu trống",
        "EnsureAccountActive → IssueTokensAsync (không bắt verify email riêng)",
    ]:
        add_bullet(doc, step)

    add_heading_styled(doc, "6.5. EmailVerificationService & PasswordResetService", 2)
    add_para(doc, "EmailVerificationService:", bold=True)
    add_bullet(doc, "SendVerificationLinkAsync: rate limit max/hour + cooldown; token URL-safe Base64; lưu hash purpose=EmailVerification; invalidate token cũ; gửi email")
    add_bullet(doc, "VerifyByTokenAsync: hash lookup; idempotent nếu đã verified; used/expired → lỗi; mark used; MarkEmailVerifiedAsync")
    add_para(doc, "PasswordResetService:", bold=True)
    add_bullet(doc, "SendResetLinkAsync: user null/inactive/no password → return null (không leak); rate limit; purpose=PasswordReset")
    add_bullet(doc, "ResetPasswordByTokenAsync: validate policy+confirm; mark used; update hash; revoke all refresh tokens")

    add_heading_styled(doc, "6.6. AuthController – API", 2)
    add_para(doc, "Route: api/auth | EnableRateLimiting(Auth)")
    add_table(
        doc,
        ["HTTP", "Path", "Auth", "Service method"],
        [
            ["POST", "register", "Public", "RegisterAsync"],
            ["POST", "verify-email", "Public", "VerifyEmailAsync"],
            ["POST", "resend-verification-email", "Public", "ResendVerificationEmailAsync"],
            ["POST", "forgot-password", "Public", "ForgotPasswordAsync"],
            ["POST", "reset-password", "Public", "ResetPasswordAsync"],
            ["POST", "change-password", "Authorize", "ChangePasswordAsync"],
            ["POST", "login", "Public", "LoginAsync"],
            ["POST", "google-login", "Public", "GoogleAuthService.LoginWithGoogleAsync"],
            ["POST", "refresh-token", "Public", "RefreshTokenAsync"],
            ["POST", "logout", "Public", "LogoutAsync"],
            ["POST", "logout-all", "Authorize", "LogoutAllAsync"],
            ["GET", "me", "Authorize", "GetMeAsync"],
            ["PUT", "me", "Authorize", "UpdateProfileAsync"],
            ["POST", "me/avatar", "Authorize", "UpdateAvatarAsync"],
        ],
    )

    # ========== 7. BOOKING ==========
    add_heading_styled(doc, "7. BACKEND – BOOKING / HOLD / PAY / CANCEL / CHECK-IN", 1)
    add_heading_styled(doc, "7.1. Interface IBookingService", 2)
    add_code(
        doc,
        "Task<BookingResponse> CreateAsync(Guid userId, CreateBookingRequest request, CancellationToken ct)\n"
        "Task<BookingResponse> HoldAsync(Guid userId, CreateBookingRequest request, CancellationToken ct)\n"
        "Task<List<CabinAvailabilityResponse>> GetCabinAvailabilityAsync(Guid scheduleId, CancellationToken ct)\n"
        "Task<List<UserBookingListItemResponse>> GetUserBookingsAsync(Guid userId, CancellationToken ct)\n"
        "Task ConfirmPaymentAsync(Guid bookingId, Guid userId, CancellationToken ct)\n"
        "Task<CancelBookingResult> CancelAsync(Guid bookingId, Guid userId, CancellationToken ct)\n"
        "Task<CheckInBookingResponse> CheckInAsync(CheckInBookingRequest request, CancellationToken ct)\n"
        "Task<int> CancelExpiredHoldsAsync(CancellationToken ct)\n"
        "Task<int> SendHoldRemindersAsync(CancellationToken ct)",
    )

    add_heading_styled(doc, "7.2. CreateAsync – Đặt tour (status = pending)", 2)
    for s in [
        "Tìm schedule + tour; không có → ScheduleNotFound",
        "Nếu boat.compliance_status blocked (Hidden/Locked) → BoatBlockedCompliance",
        "Kiểm tra user đã có booking active cùng tour cùng ngày → từ chối",
        "Nếu có cabin: load cabin của boat; quantity + đã book ≤ total_rooms",
        "Tạo booking với breakdown giá (base/cabin/service/discount/total), status=pending",
        "Thêm booking_cabin, booking_service; SaveChanges → BookingResponse",
    ]:
        add_bullet(doc, s)

    add_heading_styled(doc, "7.3. HoldAsync – Giữ chỗ (status = holding)", 2)
    for s in [
        "Schedule tồn tại + boat không blocked",
        "Xác định role agent → HoldPolicy.CalculateHoldDuration:",
        "  • Sát ngày khởi hành (NoHoldWithinDays) → cấm giữ (HoldNotAllowed)",
        "  • B2C: giữ B2CHoldMinutes (mặc định 30 phút)",
        "  • B2B: ≥30 ngày trước → 48h; ≥7 ngày → 24h; còn lại → 4h",
        "Set hold_expired_at = now + duration, status=holding",
        "Lưu cabins/services tương tự Create; trả HoldExpiredAt cho FE đếm ngược",
    ]:
        add_bullet(doc, s)

    add_heading_styled(doc, "7.4. ConfirmPaymentAsync", 2)
    for s in [
        "Tìm booking thuộc user",
        "Chỉ xử lý nếu status pending hoặc holding",
        "Holding đã quá hạn → HoldExpired",
        "Set status=confirmed, clear hold_expired_at",
        "Notification cho khách + owner (best-effort)",
        "Email SendBookingStatusEmailAsync confirmed (best-effort)",
    ]:
        add_bullet(doc, s)

    add_heading_styled(doc, "7.5. CancelAsync", 2)
    for s in [
        "Đã cancelled → lỗi",
        "Nếu status đã trả (paid/confirmed/checked_in) và còn ≥ 2 ngày trước khởi hành → RefundToWalletAsync",
        "Set cancelled + cancel_reason (refunded / no-refund / generic)",
        "Trả CancelBookingResult (Refunded, AmountRefunded)",
    ]:
        add_bullet(doc, s)

    add_heading_styled(doc, "7.6. CheckInAsync", 2)
    for s in [
        "Parse BookingCode: full GUID hoặc 8 ký tự đầu",
        "Đã check-in / cancelled / pending / completed → lỗi riêng",
        "Chỉ paid hoặc confirmed được check-in",
        "Set checked_in → trả customer/tour/boat/mã/checkedInAt",
    ]:
        add_bullet(doc, s)

    add_heading_styled(doc, "7.7. Worker helpers", 2)
    add_bullet(doc, "CancelExpiredHoldsAsync: huỷ holding có hold_expired_at ≤ now, reason \"Tự động huỷ do quá hạn giữ chỗ\"")
    add_bullet(doc, "SendHoldRemindersAsync: holds B2B sắp hết hạn trong ReminderBeforeExpiryHours; gửi email; set hold_reminder_sent=true")

    add_heading_styled(doc, "7.8. BookingsController", 2)
    add_table(
        doc,
        ["HTTP", "Path", "Auth", "Method"],
        [
            ["POST", "api/bookings", "Authorize", "CreateAsync"],
            ["POST", "api/bookings/hold", "Authorize", "HoldAsync"],
            ["GET", "api/bookings/schedules/{scheduleId}/cabins", "AllowAnonymous", "GetCabinAvailabilityAsync"],
            ["GET", "api/bookings", "Authorize", "GetUserBookingsAsync"],
            ["PUT", "api/bookings/{id}/pay", "Authorize", "ConfirmPaymentAsync"],
            ["PUT", "api/bookings/{id}/cancel", "Authorize", "CancelAsync"],
        ],
    )
    add_para(doc, "Check-in public: PUT api/public/tours/bookings/check-in → PublicToursController → CheckInAsync")

    # ========== 8. BILLING & WALLET ==========
    add_heading_styled(doc, "8. BACKEND – BILLING PAYOS & WALLET", 1)
    add_heading_styled(doc, "8.1. BillingService", 2)
    add_para(doc, "GetFinancialSummaryAsync(ownerId):", bold=True)
    add_bullet(doc, "Commission = tổng doanh thu booking × BillingOptions.Commission (vd 8%)")
    add_bullet(doc, "Maintenance = tổng giá dịch vụ bảo trì đã duyệt")
    add_bullet(doc, "Dock rental = mỗi tháng có dock_schedule × MonthlyDockRental (vd 5.000.000)")
    add_bullet(doc, "TotalPaid từ owner_payment status=paid; RemainingBalance = max(0, TotalOwed - TotalPaid)")
    add_para(doc, "InitiatePaymentAsync(ownerId):", bold=True)
    add_bullet(doc, "EnsurePayOSConfigured; remaining≤0 → lỗi")
    add_bullet(doc, "orderCode = milliseconds từ OrderCodeEpoch; PayOS PaymentRequests.CreateAsync")
    add_bullet(doc, "Lưu owner_payment pending; trả CheckoutUrl, QR, account info")
    add_para(doc, "HandlePayOSWebhookAsync(Webhook):", bold=True)
    add_bullet(doc, "Verify chữ ký PayOS; tìm payment pending theo OrderCode → paid + paid_at")
    add_bullet(doc, "SignalR BillingHub: PaymentReceived tới group owner + broadcast")
    add_bullet(doc, "Trả code 00 / 01 / 99")

    add_heading_styled(doc, "8.2. WalletService", 2)
    add_bullet(doc, "GetBalanceAsync / GetWithdrawalsAsync")
    add_bullet(doc, "RequestWithdrawAsync: trừ balance ngay; tạo wallet_withdrawal pending")
    add_para(doc, "AdminWithdrawalsService: Approve → approved; Reject → hoàn tiền lại ví + email")

    # ========== 9. OWNER REG + COMPLIANCE ==========
    add_heading_styled(doc, "9. BACKEND – OWNER REGISTRATION & BOAT COMPLIANCE", 1)
    add_heading_styled(doc, "9.1. OwnerRegistrationService.RegisterOwnerAsync", 2)
    for s in [
        "User tồn tại; chưa có owner_profile",
        "Validate entity_type (individual/company…)",
        "Transaction: tạo owner_profile Pending, is_verified=false",
        "Upload owner documents Cloudinary; national_id → license_image",
        "Mỗi vessel: upload images/docs; tạo boat Pending, compliance_status=Valid",
        "Certificates: EnsureActiveCodeAsync scope boat; upload → boat_certificate Pending",
        "Tạo boat_images; Commit; email đăng ký thành công; Rollback nếu lỗi",
    ]:
        add_bullet(doc, s)

    add_heading_styled(doc, "9.2. BoatComplianceService.RunComplianceCheckAsync", 2)
    for s in [
        "Lấy \"hôm nay\" theo timezone cấu hình (SE Asia Standard Time)",
        "MarkExpiredAsync cho certificates hết hạn",
        "Gửi reminder sắp hết hạn → MarkReminderSentAsync",
        "DetermineComplianceStatus mỗi boat:",
        "  • Cert Expired quá GracePeriodDays → Locked",
        "  • Cert Expired trong grace → Hidden",
        "  • Approved sắp hết hạn (ReminderDaysBeforeExpiry) → Warning",
        "  • Còn lại → Valid",
        "Đổi status → deactivate boat services nếu Hidden/Locked; notify (hiện NoOpBoatComplianceNotifier)",
    ]:
        add_bullet(doc, s)

    add_heading_styled(doc, "9.3. BoatCertificateService (tóm tắt hàm)", 2)
    for x in [
        "GetByBoatIdForOwnerAsync / GetByOwnerIdAsync – kiểm tra ownership",
        "UploadAsync – không deprecated type; không trùng type; Cloudinary; Pending",
        "RenewAsync – xóa file cũ Cloudinary; reset Pending",
        "GetPending/Approved/ExpiringForAdminAsync",
        "ApproveAsync / RejectAsync (reason bắt buộc) + notify",
        "UnlockBoatAsync – chỉ Hidden/Locked; cần ≥1 cert Approved còn hạn → Valid",
    ]:
        add_bullet(doc, x)

    # ========== 10. ADMIN ==========
    add_heading_styled(doc, "10. BACKEND – ADMIN SERVICES", 1)
    add_heading_styled(doc, "10.1. AdminOwnersService", 2)
    add_bullet(doc, "GetVerificationsAsync: liệt kê owner_profile + số boat + docs + vessels + certificates")
    add_bullet(doc, "ApproveVerificationAsync: Transaction status=Verified, is_verified=true; đảm bảo role owner; boat Pending→Idle; email duyệt")
    add_bullet(doc, "RejectVerificationAsync: status=Rejected; boat Pending→Rejected")

    add_heading_styled(doc, "10.2. AdminUserService", 2)
    add_bullet(doc, "GetUsersAsync (paging), GetStatsAsync, GetByIdAsync")
    add_bullet(doc, "UpdateAsync: không sửa chính mình; cập nhật name/phone/isActive; deactivate → revoke tokens")
    add_bullet(doc, "UpdateRolesAsync: roles cho phép admin/owner/user/agent")
    add_bullet(doc, "DeactivateAsync: is_active=false + revoke tokens")

    add_heading_styled(doc, "10.3. Các admin service khác", 2)
    add_bullet(doc, "AdminDashboardService: GetStatsAsync, GetTopToursAsync")
    add_bullet(doc, "AdminMaintenancesService: pending/all; Approve/Reject + email")
    add_bullet(doc, "AdminPromotionsService: CRUD + Approve/Reject/ToggleActive; code unique uppercase")
    add_bullet(doc, "AdminWithdrawalsService: Approve/Reject (reject hoàn ví)")

    # ========== 11. OTHER SERVICES ==========
    add_heading_styled(doc, "11. BACKEND – CÁC SERVICE CÒN LẠI (HÀM + LOGIC)", 1)

    add_heading_styled(doc, "11.1. ChatService", 2)
    add_bullet(doc, "GetConversationsAsync / GetMessagesAsync (phải là member)")
    add_bullet(doc, "StartConversationAsync(bookingId): chỉ khách hoặc owner tour; idempotent; thêm 2 members")
    add_bullet(doc, "SendMessageAsync: lưu message; SignalR ReceiveMessage; email nếu user offline (ChatHub.IsUserOnline)")
    add_bullet(doc, "MarkAsReadAsync")

    add_heading_styled(doc, "11.2. OwnerToursDashboardService", 2)
    add_bullet(doc, "GetStatsAsync, GetSchedulesAsync(month,year), GetRecentBookingsAsync, GetResourcesAsync")
    add_bullet(doc, "CreateScheduleAsync: ownership, overlap tour/boat, duration mặc định 120 phút")
    add_bullet(doc, "UpdateBookingStatusAsync: nếu cancel đơn đã trả → hoàn ví + email")
    add_bullet(doc, "Seed helpers: CleanSeedDataAsync, RenameBoatsAsync, SeedToursAsync")

    add_heading_styled(doc, "11.3. Tour / Public catalog", 2)
    add_bullet(doc, "PublicTourSearchService: SearchAsync (paging filters), GetPopularDestinationsAsync")
    add_bullet(doc, "PublicTourCatalogService: GetActiveTourAsync, GetTourImagesAsync, GetTourFaqsAsync")
    add_bullet(doc, "TourService / TourScheduleService / RouteService / TourContentService / FaqService / TourImageService: CRUD owner/legacy")
    add_bullet(doc, "OwnerServicesRegistrationService: tạo tour pending + cabins/combos/faqs/routes/default schedule")

    add_heading_styled(doc, "11.4. Boat / Dock / Review / Wishlist / Notification", 2)
    add_bullet(doc, "BoatService: Admin CRUD + Owner-scoped CRUD (GetBoatsByOwnerAsync, CreateByOwnerAsync…)")
    add_bullet(doc, "BoatCabinService / BoatAddonService / BoatImageService / BoatMaintenanceService")
    add_bullet(doc, "DockService / DockScheduleService: CRUD dock + gán lịch neo; email gán dock")
    add_bullet(doc, "ReviewService: paginated; can-review (unreviewed completed bookings); create multipart Cloudinary; update/delete own")
    add_bullet(doc, "WishlistService: get; toggle; get wishlisted IDs")
    add_bullet(doc, "NotificationService: create + recipients; push SignalR qua ChatHub; mark read / mark all")
    add_bullet(doc, "CloudinaryService: Upload/Delete image & video")
    add_bullet(doc, "EmailSender: nhiều template (verify, reset, hold reminder, booking status, owner reg, maintenance, withdrawal, chat, schedule, dock…)")
    add_bullet(doc, "SystemService / CertificateTypeService / MaintenanceServicesService / OwnerDocumentService")

    # ========== 12. CONTROLLERS ==========
    add_heading_styled(doc, "12. BACKEND – CONTROLLERS & API ENDPOINTS", 1)
    add_table(
        doc,
        ["Controller", "Base route", "Auth", "Chức năng chính"],
        [
            ["AuthController", "api/auth", "Public + Authorize", "Đăng ký/đăng nhập/token/profile"],
            ["BookingsController", "api/bookings", "Authorize", "Create/Hold/Pay/Cancel/List"],
            ["PublicToursController", "api/public/tours", "Public", "Catalog + check-in"],
            ["TourSearchController", "api/tours/search", "Public", "Search + popular destinations"],
            ["OwnerBillingController", "api/owner/billing", "owner / webhook public", "Summary, PayOS pay, webhook"],
            ["OwnerRegistrationController", "api/owner", "Authorize", "POST register multipart"],
            ["OwnerDocumentsController", "api/owner/documents", "Authorize", "List/Upload documents"],
            ["OwnerBoatsController", "api/owner/boats", "owner", "CRUD boats, images, maint, certs"],
            ["OwnerToursDashboardController", "api/owner/tours-dashboard", "owner", "Stats/schedules/bookings"],
            ["OwnerServicesController", "api/owner/services", "owner", "Đăng ký dịch vụ tour"],
            ["OwnerPromotionsController", "api/owner/promotions", "owner", "KM của owner"],
            ["AdminOwnersController", "api/admin/owners", "admin", "Duyệt verification"],
            ["AdminUsersController", "api/admin/users", "admin", "CRUD users/roles"],
            ["AdminBoatCertificatesController", "api/admin/...", "admin", "Duyệt GCN + unlock"],
            ["AdminCertificateTypesController", "api/admin/certificate-types", "admin", "CRUD loại GCN"],
            ["AdminDashboardController", "api/admin/dashboard", "admin", "Stats, top-tours"],
            ["AdminMaintenancesController", "api/admin/maintenances", "admin", "Duyệt bảo trì"],
            ["AdminWithdrawalsController", "api/admin/withdrawals", "admin", "Duyệt rút tiền"],
            ["AdminPromotionsController", "api/admin/promotions", "admin", "CRUD + approve KM"],
            ["BoatsController", "api/admin/boats + api/boats", "admin/public", "Admin CRUD; public list"],
            ["BoatCabins/Addons/Images", "api/admin/boats/{id}/...", "admin", "Nested resources"],
            ["DocksController", "api/docks / api/admin/docks", "admin/owner", "CRUD dock + schedules"],
            ["ChatController", "api/chat", "Authorize", "Conversations/messages"],
            ["NotificationsController", "api/Notifications", "Authorize", "List/read"],
            ["WalletController", "api/wallet", "Authorize", "Balance/withdraw"],
            ["WishlistsController", "api/Wishlists", "Authorize", "Toggle wishlist"],
            ["ReviewsController", "api/Reviews", "Public GET / Auth CUD", "Reviews"],
            ["TourContent/Faqs", "api/tours/{id}/...", "owner", "Nội dung tour"],
            ["SystemController", "api/system", "Public", "boat-types, certificate-types"],
            ["Legacy Tour/Route/Schedule", "api/legacy/... , api/tour-schedules", "tùy", "CRUD legacy"],
        ],
    )

    # ========== 13. JOBS & HUBS ==========
    add_heading_styled(doc, "13. BACKEND – BACKGROUND JOBS & SIGNALR HUBS", 1)
    add_heading_styled(doc, "13.1. BoatComplianceBackgroundService", 2)
    add_para(
        doc,
        "PeriodicTimer theo BoatComplianceOptions.CheckIntervalHours (vd 6h). Chạy ngay khi start rồi lặp. "
        "Scope DI → IBoatComplianceService.RunComplianceCheckAsync. Catch exception không crash worker."
    )
    add_heading_styled(doc, "13.2. SeatHoldCleanupBackgroundService", 2)
    add_para(
        doc,
        "Interval BookingHoldOptions.CleanupIntervalMinutes (mặc định 1 phút). "
        "Gọi CancelExpiredHoldsAsync rồi SendHoldRemindersAsync. Log số booking huỷ / email nhắc."
    )
    add_heading_styled(doc, "13.3. ChatHub [Authorize]", 2)
    add_bullet(doc, "ConcurrentDictionary online users (userId → connectionIds)")
    add_bullet(doc, "OnConnectedAsync: add connection + Groups.AddToGroup(userId)")
    add_bullet(doc, "OnDisconnectedAsync: remove connection")
    add_bullet(doc, "IsUserOnline(userId) dùng bởi ChatService để quyết định gửi email")
    add_heading_styled(doc, "13.4. BillingHub", 2)
    add_para(doc, "Hub trống; BillingService dùng IHubContext<BillingHub> gửi event PaymentReceived.")

    doc.add_page_break()

    # ========== 14. FE CUSTOMER STRUCTURE ==========
    add_heading_styled(doc, "14. FRONTEND CUSTOMER/OWNER – CẤU TRÚC & ROUTES", 1)
    add_para(doc, "src/ gồm: assets, components (layouts/ui/routes), config, constants, contexts, data, hooks, i18n, interfaces, lib, pages, services, styles.")
    add_heading_styled(doc, "14.1. Routes chính (App.tsx)", 2)
    add_para(doc, "AuthLayout:", bold=True)
    add_bullet(doc, "/sign-in, /sign-up, /verify-email (+pending/success), /forgot-password, /reset-password (+success)")
    add_para(doc, "MainLayout – Public:", bold=True)
    add_bullet(doc, "/, /tours, /tours/:id, /tours/:id/booking, /boats/:boatId, /become-owner")
    add_para(doc, "MainLayout – ProtectedRoute (đã login):", bold=True)
    add_bullet(doc, "/owner-registration, /profile, /my-tours, /wishlist, /wallet, /inbox")
    add_para(doc, "Owner – ProtectedRoute roles=['owner'] + OwnerLayout:", bold=True)
    add_bullet(doc, "/owner, /owner/boats, /owner/boats/new, /owner/boats/:boatId/edit, /owner/admin-maintenance, /owner/dock-map, /owner/services, /owner/tours, /owner/revenue, /owner/certificates, /owner/documents")
    add_para(doc, "Catch-all * → home. Pages dùng React.lazy + Suspense.")

    add_heading_styled(doc, "14.2. ProtectedRoute logic", 2)
    add_code(
        doc,
        "if (!isAuthenticated) Navigate → /sign-in (giữ state.from)\n"
        "if (roles && !user.roles.some(r => roles.includes(r))) Navigate → /\n"
        "else <Outlet />",
    )

    # ========== 15. FE AUTH AXIOS ==========
    add_heading_styled(doc, "15. FRONTEND CUSTOMER – AUTH, AXIOS, TOKEN REFRESH", 1)
    add_heading_styled(doc, "15.1. AuthContext", 2)
    add_bullet(doc, "State: token (access_token), user (normalize từ localStorage), isAuthenticated=!!token")
    add_bullet(doc, "normalizeUser: map name/fullName, email, roles∈{user,owner,admin}, avatar, phone, address, hasOwnerProfile")
    add_bullet(doc, "login(token,user): ghi localStorage + set state")
    add_bullet(doc, "logout(): xóa access/refresh/user")
    add_bullet(doc, "reloadUser(): GET /auth/me")
    add_bullet(doc, "Mount: nếu có refresh_token → AuthServices.refreshToken → cập nhật tokens; rồi getProfile đồng bộ role (vd sau khi admin duyệt owner)")
    add_bullet(doc, "Đồng bộ đa tab qua storage event")

    add_heading_styled(doc, "15.2. Axios layers", 2)
    add_para(doc, "services/axios.ts – Api (có Bearer) + Axios (không Bearer cho login/register):", bold=True)
    add_bullet(doc, "Request: Authorization Bearer access_token")
    add_bullet(doc, "503 → /maintenance")
    add_bullet(doc, "code 1201 ACCOUNT_INACTIVE → dialog/toast, clear session, redirect sign-in")
    add_bullet(doc, "401: refreshAccessTokenShared(); retry 1 lần; fail → clear + redirect")
    add_para(doc, "services/api.ts – client dùng bởi hầu hết domain services:", bold=True)
    add_bullet(doc, "Bearer + Accept-Language từ i18nextLng; 401/403 shared refresh + retry")
    add_para(doc, "auth-token-refresh.ts:", bold=True)
    add_bullet(doc, "refreshAccessTokenShared(): singleton Promise tránh gọi song song (chống reuse detection)")
    add_bullet(doc, "POST /auth/refresh-token {refreshToken}; code===1000 → lưu tokens")
    add_bullet(doc, "clearAuthSession(): xóa access, refresh, user")

    add_heading_styled(doc, "15.3. Luồng đăng nhập trang sign-in", 2)
    for s in [
        "AuthServices.login hoặc googleLogin",
        "unwrapEnvelope lấy token + refreshToken",
        "saveAuthTokens + getProfile → mapProfileToUser → AuthContext.login",
        "Navigate from state hoặc home",
        "Nếu code EMAIL_NOT_VERIFIED (1204) → /verify-email/pending",
    ]:
        add_bullet(doc, s)

    # ========== 16. FE SERVICES ==========
    add_heading_styled(doc, "16. FRONTEND CUSTOMER – SERVICES (HÀM + ENDPOINT)", 1)
    add_para(doc, "Envelope chuẩn: { code: number, result: T, message? }. Success thường code = 1000.")

    add_heading_styled(doc, "16.1. AuthServices", 2)
    add_table(
        doc,
        ["Hàm FE", "HTTP Endpoint"],
        [
            ["login / register", "POST /auth/login | /auth/register"],
            ["verifyEmail / resendVerificationEmail", "POST /auth/verify-email | resend-verification-email"],
            ["forgotPassword / resetPassword", "POST /auth/forgot-password | reset-password"],
            ["changePassword", "POST /auth/change-password"],
            ["googleLogin", "POST /auth/google-login"],
            ["refreshToken", "POST /auth/refresh-token"],
            ["getProfile / updateProfile / updateAvatar", "GET|PUT /auth/me ; POST /auth/me/avatar"],
            ["logout / logoutAll", "POST /auth/logout | logout-all"],
            ["registerOwner", "POST /owner/register (multipart)"],
        ],
    )

    add_heading_styled(doc, "16.2. tourService", 2)
    add_table(
        doc,
        ["Hàm", "Endpoint"],
        [
            ["searchTours", "GET /tours/search"],
            ["getPopularDestinations", "GET /public/tours/destinations/popular"],
            ["getPublicTourById / getTourImages / getTourFaqs", "GET /public/tours/{id}[/images|/faqs]"],
            ["getTourSchedules", "GET /tour-schedules/tour/{id}"],
            ["getToursDashboardStats/Schedules/RecentBookings/Resources", "GET /owner/tours-dashboard/..."],
            ["createTourSchedule", "POST /owner/tours-dashboard/schedule"],
            ["updateBookingStatus", "POST /owner/tours-dashboard/bookings/{id}/status"],
        ],
    )

    add_heading_styled(doc, "16.3. bookingService", 2)
    add_table(
        doc,
        ["Hàm", "Endpoint", "Ghi chú"],
        [
            ["createBooking", "POST /bookings", "status pending"],
            ["holdBooking", "POST /bookings/hold", "trả holdExpiredAt đếm ngược"],
            ["getUserBookings", "GET /bookings", "My Tours + QR"],
            ["getCabinAvailability", "GET /bookings/schedules/{id}/cabins", "public"],
            ["confirmPayment", "PUT /bookings/{id}/pay", "xác nhận TT"],
            ["cancelBooking", "PUT /bookings/{id}/cancel", "có thể refund ví"],
        ],
    )

    add_heading_styled(doc, "16.4. boatService & liên quan", 2)
    add_bullet(doc, "getAllPublic, getByIdPublic, getOwnerBoats/stats/byId, create/update/deleteByOwner, upload/deleteBoatImage")
    add_bullet(doc, "registerPortMaintenances, deleteOwnerMaintenance; admin approve/reject maintenances")
    add_bullet(doc, "boatServiceApi: CRUD /admin/boats/{id}/services (+ toggle)")
    add_bullet(doc, "cabinService: CRUD /admin/boats/{id}/cabins")

    add_heading_styled(doc, "16.5. Chat + SignalR", 2)
    add_bullet(doc, "chatService: getConversations, getMessages, startConversation(bookingId), sendMessage, markAsRead")
    add_bullet(doc, "chatSignalRService: Hub {VITE_API_URL}/hub/chat; accessTokenFactory; withAutomaticReconnect")
    add_bullet(doc, "Events: ReceiveMessage, ReceiveNotification; startConnection / stopConnection")

    add_heading_styled(doc, "16.6. Các service khác", 2)
    add_bullet(doc, "walletService: getBalance, getWithdrawals, requestWithdraw")
    add_bullet(doc, "billingService: getFinancialSummary, initiatePayment (PayOS)")
    add_bullet(doc, "certificateService / ownerDocumentService: types, upload, renew, list docs")
    add_bullet(doc, "wishlistService: getWishlists, getWishlistedTourIds, toggleWishlist")
    add_bullet(doc, "reviewService: getReviewsByTourId, canReviewTour, create/update/delete multipart")
    add_bullet(doc, "notificationService: getNotifications, markAsRead, markAllAsRead")
    add_bullet(doc, "dockService: CRUD docks + schedules")
    add_bullet(doc, "weatherService: Open-Meteo geocoding + forecast (không qua backend)")
    add_bullet(doc, "system-service: getBoatTypes → GET /system/boat-types")

    # ========== 17. FE PAGES ==========
    add_heading_styled(doc, "17. FRONTEND CUSTOMER – PAGES & LUỒNG UI", 1)
    add_table(
        doc,
        ["Trang", "Chức năng", "Services chính"],
        [
            ["auth/*", "Login/Register/Verify/Reset", "AuthServices"],
            ["home", "Landing featured tours", "tourService, wishlistService"],
            ["tours/tour-list", "Tìm/filter tour", "searchTours, toggleWishlist"],
            ["tours/tour-detail", "Chi tiết + reviews", "catalog + reviewService"],
            ["tours/booking", "Wizard đặt chỗ", "hold/create/pay + cabin availability"],
            ["boats/boat-detail", "Chi tiết thuyền public", "boatService.getByIdPublic"],
            ["owner-registration", "Form multipart đăng ký owner", "registerOwner, cert types, boat types"],
            ["profile", "Hồ sơ/avatar/đổi MK", "updateProfile/Avatar, changePassword"],
            ["my-tours", "Booking + QR + start chat", "getUserBookings, cancel, chatService"],
            ["wishlist / wallet / inbox", "Yêu thích / ví / chat RT", "wishlist / wallet / chat+SignalR"],
            ["owner/index", "Dashboard owner", "getOwnerStats, boats, certs"],
            ["owner/boats + boat-form", "CRUD thuyền/ảnh/cert/DV", "boatService, certificateService"],
            ["owner/tours", "Lịch & booking dashboard", "tours-dashboard APIs"],
            ["owner/revenue", "Tài chính + PayOS", "billingService"],
            ["owner/certificates/documents", "GCN & hồ sơ pháp lý", "certificate/document services"],
            ["owner/dock-map / services", "Bản đồ bến / bảo trì cảng", "dockService, maintenance services"],
        ],
    )

    doc.add_page_break()

    # ========== 18. FE ADMIN ==========
    add_heading_styled(doc, "18. FRONTEND ADMIN / KIOSK – CHI TIẾT", 1)
    add_heading_styled(doc, "18.1. Khác biệt so với customer FE", 2)
    add_table(
        doc,
        ["Tiêu chí", "Customer FE", "Admin FE"],
        [
            ["Đối tượng", "Customer + Owner", "Admin + Owner nhẹ + Kiosk"],
            ["Google OAuth", "Có", "Không"],
            ["SignalR", "Có /hub/chat", "Không"],
            ["Silent refresh token", "Có", "Không"],
            ["Refresh on mount", "Có", "Không"],
            ["Default i18n", "en", "vn"],
            ["Kiosk QR", "Không", "html5-qrcode"],
        ],
    )

    add_heading_styled(doc, "18.2. Routes Admin", 2)
    add_bullet(doc, "Auth: /admin/sign-in, /admin/sign-up")
    add_bullet(doc, "Kiosk: /kiosk-checkin (KHÔNG ProtectedRoute – public fullscreen)")
    add_bullet(doc, "Admin (roles=['admin']): /admin, users, owner-verification, docks, promotions, revenue, top-tours, tour-approvals, boats, reviews, faqs, notifications, audit-logs, approvals, legal-compliance")
    add_bullet(doc, "Owner (roles=['owner']): /owner/boats, /owner/promotions; /owner → boats")
    add_bullet(doc, "Catch-all → /admin")

    add_heading_styled(doc, "18.3. AuthContext Admin", 2)
    add_para(
        doc,
        "Đơn giản hơn: keys token/user/access_token; login lưu JSON.stringify(token) vào access_token; "
        "không refresh-on-mount, không reloadUser. 401 interceptor clear session → /admin/sign-in (không silent refresh)."
    )

    add_heading_styled(doc, "18.4. Services Admin", 2)
    add_table(
        doc,
        ["Service", "Hàm / Endpoint chính"],
        [
            ["AuthServices", "login, register, forgot/reset, refreshToken, getProfile, changePassword, logout"],
            ["approvalsApi", "GET/POST /admin/maintenances; /admin/withdrawals (+approve/reject)"],
            ["boatApi", "CRUD /admin/boats + cabins/services/images"],
            ["certificateApi", "pending/approved/expiring; approve/reject; unlockBoat; certificate-types CRUD"],
            ["dockApi", "CRUD /admin/docks + schedules"],
            ["promotionsApi", "Admin full CRUD+approve; Owner get/create/delete"],
            ["tourApprovalApi", "GET /legacy/tours; PUT status active/rejected"],
            ["checkInService.checkIn", "PUT /public/tours/bookings/check-in {bookingCode}"],
        ],
    )

    add_heading_styled(doc, "18.5. Pages Admin & trạng thái tích hợp API", 2)
    add_table(
        doc,
        ["Trang", "API thật?", "Ghi chú"],
        [
            ["admin/index", "Có", "GET /admin/dashboard/stats"],
            ["admin/users", "Có", "CRUD users + roles"],
            ["admin/owner-verification", "Có", "/admin/owners/verifications"],
            ["admin/boats", "Có", "boatApi + unlock"],
            ["admin/docks", "Có", "dockApi"],
            ["admin/promotions", "Có", "promotionsApi"],
            ["admin/tour-approvals", "Có", "legacy tours"],
            ["admin/approvals", "Có", "maintenances + withdrawals"],
            ["admin/legal-compliance", "Có", "certificateApi"],
            ["admin/top-tours", "Có", "/admin/dashboard/top-tours"],
            ["admin/revenue, reviews, faqs, notifications, audit-logs", "Mock UI", "data cứng / chưa nối API"],
            ["kiosk-checkin", "Có", "checkInService + html5-qrcode"],
            ["owner/promotions", "Có", "owner promotions API"],
            ["owner/boats (admin FE)", "Mock", "local state, không boat-api"],
        ],
    )

    add_heading_styled(doc, "18.6. Kiosk check-in – logic UI", 2)
    for s in [
        "Màn hình fullscreen; có thể nhập mã thủ công hoặc quét QR (html5-qrcode)",
        "Gọi checkInService.checkIn(bookingCode)",
        "Validate HTTP 200 && code===1000 → hiển thị CheckInBookingResponse (tên KH, tour, thuyền, số người, giờ xuất bến, checkedInAt)",
        "Lỗi → toast message từ backend",
    ]:
        add_bullet(doc, s)

    # ========== 19. E2E FLOWS ==========
    add_heading_styled(doc, "19. LUỒNG NGHIỆP VỤ END-TO-END", 1)

    add_heading_styled(doc, "19.1. Đăng ký & đăng nhập khách", 2)
    add_para(
        doc,
        "Register → email verification link → VerifyEmail → Login (hoặc Google auto-verify) → "
        "JWT access + refresh (hash SHA256) → FE lưu localStorage → gắn Bearer mọi request."
    )

    add_heading_styled(doc, "19.2. Trở thành chủ thuyền", 2)
    add_para(
        doc,
        "User login → form /owner-registration multipart → OwnerRegistrationService tạo profile Pending + boats Pending + certificates Pending → "
        "Admin duyệt (ApproveVerification) → gán role owner + boats Idle → FE refresh token/me để nhận role mới → vào /owner."
    )

    add_heading_styled(doc, "19.3. Đặt tour & thanh toán", 2)
    add_para(
        doc,
        "Search public → chi tiết tour → chọn schedule → xem cabin availability → Hold (countdown) hoặc Create → "
        "ConfirmPayment → status confirmed + notify/email. Worker mỗi phút huỷ hold hết hạn. "
        "My Tours hiển thị QR (8 ký tự / GUID) để kiosk check-in."
    )

    add_heading_styled(doc, "19.4. Hủy & hoàn ví", 2)
    add_para(
        doc,
        "CancelAsync: nếu đã trả và ≥2 ngày trước khởi hành → cộng tiền vào user_wallet. "
        "Owner update booking status cancel cũng có thể hoàn ví. User rút tiền → Admin approve/reject (reject hoàn lại số dư)."
    )

    add_heading_styled(doc, "19.5. Billing chủ thuyền", 2)
    add_para(
        doc,
        "Owner xem financial-summary (commission + maintenance + dock rental − đã trả) → InitiatePayment PayOS → "
        "redirect/QR → webhook verify → owner_payment paid → SignalR PaymentReceived realtime trên FE revenue."
    )

    add_heading_styled(doc, "19.6. Compliance giấy tờ", 2)
    add_para(
        doc,
        "Owner upload/renew certificates → Admin approve/reject. Job định kỳ đánh dấu Expired, Warning/Hidden/Locked. "
        "Boat Hidden/Locked bị chặn booking. Admin unlock khi có cert Approved còn hạn."
    )

    add_heading_styled(doc, "19.7. Chat theo booking", 2)
    add_para(
        doc,
        "Từ My Tours: startConversation(bookingId) → Inbox kết nối SignalR /hub/chat → send/receive realtime; "
        "nếu đối phương offline có thể gửi email."
    )

    # ========== 20. API ENVELOPE ==========
    add_heading_styled(doc, "20. ENVELOPE API, MÃ LỖI, GHI CHÚ KỸ THUẬT", 1)
    add_heading_styled(doc, "20.1. Envelope thành công", 2)
    add_code(doc, "{ \"code\": 1000, \"result\": { ... }, \"message\": \"...\" }")
    add_para(
        doc,
        "Một số endpoint owner dashboard dùng shape { isSuccess, result } thay vì ApiResponse – FE cần xử lý cả hai."
    )

    add_heading_styled(doc, "20.2. Mã lỗi nổi bật (customer FE ApiErrorCode)", 2)
    add_table(
        doc,
        ["Code", "Ý nghĩa"],
        [
            ["1000", "SUCCESS"],
            ["1201", "ACCOUNT_INACTIVE / USER_DISABLED"],
            ["1202", "INVALID_CREDENTIALS"],
            ["1204", "EMAIL_NOT_VERIFIED"],
            ["14xx", "TOKEN / REFRESH_TOKEN errors, UNAUTHORIZED 1401, FORBIDDEN 1403"],
            ["429", "Auth rate limited / OTP rate limited"],
        ],
    )

    add_heading_styled(doc, "20.3. Cấu hình quan trọng (appsettings)", 2)
    add_bullet(doc, "Cors.AllowedOrigins: localhost:5173/5174 + host deploy")
    add_bullet(doc, "Jwt: issuer ddms-backend, audience ddms-client, accessTokenMinutes, refreshTokenDays")
    add_bullet(doc, "Billing: Commission 0.08, MonthlyDockRental 5000000, PayOS return/cancel URLs")
    add_bullet(doc, "BookingHold: NoHoldWithinDays, B2CHoldMinutes, B2B tiers hours, CleanupIntervalMinutes")
    add_bullet(doc, "BoatCompliance: CheckIntervalHours, ReminderDaysBeforeExpiry, GracePeriodDays, TimeZoneId")
    add_bullet(doc, "Cloudinary / PayOS / Google / Email / ConnectionStrings")

    add_heading_styled(doc, "20.4. Ghi chú rủi ro / kỹ thuật", 2)
    for x in [
        "Secrets (DB password, PayOS, Cloudinary) đang hard-code trong appsettings – cần chuyển sang User Secrets / env khi production",
        "Program.cs seed admin password rõ ràng + migrate/seed mỗi startup – cần kiểm soát production",
        "GlobalExceptionMiddleware có thể lộ stack trace ở lỗi uncategorized",
        "Admin FE: một số trang (reviews, faqs, revenue, notifications, audit-logs) còn mock UI",
        "Admin FE axios không silent-refresh; AuthContext keys token vs access_token cần đồng bộ cẩn thận",
        "OwnerToursDashboardController Wrap trả {isSuccess} khác ApiResponse chuẩn",
        "Thư mục admin thực tế đặt tên ddms-frontedn-admin (typo \"frontedn\")",
    ]:
        add_bullet(doc, x)

    add_heading_styled(doc, "20.5. DependencyInjection – nhóm đăng ký chính", 2)
    add_para(
        doc,
        "AddProjectDependencies đăng ký scoped: Tour/Schedule/Route/Content/Search, Billing, Booking, Wallet, "
        "OwnerToursDashboard, OwnerServicesRegistration, AdminDashboard/Owners/Maintenances/Withdrawals/Promotions, "
        "BoatCertificate/Compliance/CertificateType/OwnerDocument, System, MaintenanceServices, … "
        "Program.cs còn đăng ký Auth, Boat*, Dock*, Chat, Notification, PublicTour*, Review, Wishlist, Cloudinary, v.v."
    )

    # Footer
    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end.add_run(
        "— Hết tài liệu —\n"
        "Nguồn: phân tích trực tiếp source code DDMS.Backend, ddms-frontend, ddms-frontend-admin\n"
        "Ngày tạo file: 29/07/2026"
    )
    set_run_font(r, size=10, color=RGBColor(0x66, 0x66, 0x66))

    doc.save(OUT)
    print(f"Saved: {OUT}")
    return OUT


if __name__ == "__main__":
    build()
